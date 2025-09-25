import os
import time
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv

# Core dependencies
from pinecone import Pinecone, ServerlessSpec
from langchain_pinecone import PineconeVectorStore
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# File processing dependencies
import PyPDF2
import docx
import markdown
from bs4 import BeautifulSoup

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

load_dotenv()

class MultiFormatContentIngestion:
    """multi-format document ingestion pipeline for DataCamp course content"""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.md'}
    
    def __init__(self):
        self.pc = None
        self.index = None
        self.embeddings = None
        self.vector_store = None
        
    def initialize_pinecone(self):
        """Initialize Pinecone with enhanced error handling and monitoring"""
        try:
            api_key = os.environ.get("PINECONE_API_KEY")
            if not api_key:
                raise ValueError("PINECONE_API_KEY environment variable required")
                
            self.pc = Pinecone(api_key=api_key)
            index_name = os.environ.get("PINECONE_INDEX_NAME", "datacamp-courses-index")
            
            existing_indexes = [idx["name"] for idx in self.pc.list_indexes()]
            
            if index_name not in existing_indexes:
                logger.info(f"Creating Pinecone index: {index_name}")
                self.pc.create_index(
                    name=index_name,
                    dimension=384,  # sentence-transformers/all-MiniLM-L6-v2
                    metric="cosine",
                    spec=ServerlessSpec(cloud="aws", region="us-east-1"),
                )
                
                while not self.pc.describe_index(index_name).status["ready"]:
                    logger.info("Waiting for index initialization...")
                    time.sleep(3)
                    
            self.index = self.pc.Index(index_name)
            logger.info(f"Connected to Pinecone index: {index_name}")
            return True
            
        except Exception as e:
            logger.error(f"Pinecone initialization failed: {e}")
            return False

    def initialize_embeddings(self):
        """Initialize optimized embeddings model"""
        try:
            logger.info("Loading embeddings model...")
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
            logger.info("Embeddings model loaded successfully")
            return True
        except Exception as e:
            logger.error(f"Embeddings initialization failed: {e}")
            return False

    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return ""

    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return ""

    def extract_text_from_markdown(self, file_path: Path) -> str:
        """Extract text from Markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert markdown to HTML, then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
            
            # Clean up extra whitespace
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            return '\n'.join(lines)
        except Exception as e:
            logger.error(f"Error extracting text from Markdown {file_path}: {e}")
            return ""

    def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            return ""

    def extract_text_from_file(self, file_path: Path) -> str:
        """Extract text from any supported file format"""
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif extension == '.md':
            return self.extract_text_from_markdown(file_path)
        elif extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file format: {extension}")
            return ""

    def get_file_type_metadata(self, file_path: Path) -> Dict[str, str]:
        """Get file type specific metadata"""
        extension = file_path.suffix.lower()
        
        metadata = {
            'file_format': extension.replace('.', '').upper(),
            'original_filename': file_path.name
        }
        
        # Add format-specific metadata
        if extension == '.pdf':
            metadata['document_type'] = 'PDF Document'
        elif extension == '.docx':
            metadata['document_type'] = 'Word Document'
        elif extension == '.md':
            metadata['document_type'] = 'Markdown Document'
        elif extension == '.txt':
            metadata['document_type'] = 'Text Document'
            
        return metadata

    def load_course_documents(self, data_directory="data"):
        """Load and validate multi-format course files"""
        documents = []
        data_path = Path(data_directory)
        
        if not data_path.exists():
            logger.error(f"Data directory '{data_directory}' not found")
            return documents
        
        # Find all supported files
        supported_files = []
        for extension in self.SUPPORTED_EXTENSIONS:
            supported_files.extend(list(data_path.glob(f"*{extension}")))
        
        if not supported_files:
            logger.warning(f"No supported files found in {data_directory}")
            logger.info(f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}")
            return documents
        
        logger.info(f"Processing {len(supported_files)} course files...")
        logger.info(f"File formats found: {', '.join(set(f.suffix for f in supported_files))}")
        
        for file_path in supported_files:
            try:
                # Extract text based on file format
                content = self.extract_text_from_file(file_path)
                    
                if len(content) < 100:
                    logger.warning(f"Skipping {file_path.name} - insufficient content ({len(content)} chars)")
                    continue
                
                # Extract course information from filename/content
                course_info = self.extract_course_metadata(file_path, content)
                
                # Get file format metadata
                format_metadata = self.get_file_type_metadata(file_path)
                
                # Create document with enhanced metadata
                metadata = {
                    "source": file_path.name,
                    "file_path": str(file_path),
                    "content_length": len(content),
                    "course_title": course_info.get("title", "Unknown Course"),
                    "difficulty": course_info.get("difficulty", "Unknown"),
                    "language": course_info.get("language", "Unknown"),
                    "ingested_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                    **format_metadata  # Add file format metadata
                }
                
                doc = Document(
                    page_content=content,
                    metadata=metadata
                )
                
                documents.append(doc)
                logger.info(f"✅ Loaded: {file_path.name} ({format_metadata['file_format']}, {len(content):,} chars)")
                
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                continue
        
        logger.info(f"Successfully loaded {len(documents)} course documents")
        
        # Log file format distribution
        format_counts = {}
        for doc in documents:
            fmt = doc.metadata.get('file_format', 'UNKNOWN')
            format_counts[fmt] = format_counts.get(fmt, 0) + 1
        
        logger.info("File format distribution:")
        for fmt, count in format_counts.items():
            logger.info(f"  • {fmt}: {count} files")
        
        return documents

    def extract_course_metadata(self, file_path: Path, content: str):
        """Extract course metadata from filename and content"""
        metadata = {}
        filename = file_path.stem.lower()
        content_lower = content.lower()
        
        # Extract difficulty level
        difficulty_indicators = ['beginner', 'intermediate', 'advanced', 'basic', 'intro', 'fundamental']
        for level in difficulty_indicators:
            if level in filename or level in content_lower[:500]:
                if level in ['basic', 'intro', 'fundamental', 'beginner']:
                    metadata["difficulty"] = "Beginner"
                elif level == 'intermediate':
                    metadata["difficulty"] = "Intermediate"
                elif level == 'advanced':
                    metadata["difficulty"] = "Advanced"
                break
        
        # Extract programming language with expanded detection
        languages = {
            'python': ['python', 'pandas', 'numpy', 'matplotlib', 'scikit', 'jupyter'],
            'r': ['r programming', ' r ', 'ggplot', 'dplyr', 'tidyverse'],
            'sql': ['sql', 'database', 'query', 'select', 'join'],
            'scala': ['scala', 'spark'],
            'julia': ['julia'],
            'bash': ['bash', 'shell', 'command line'],
            'git': ['git', 'github', 'version control'],
            'javascript': ['javascript', 'js', 'node'],
            'java': ['java'],
            'c++': ['c++', 'cpp']
        }
        
        for lang, keywords in languages.items():
            if any(keyword in filename or keyword in content_lower[:1000] for keyword in keywords):
                metadata["language"] = lang.upper() if lang == 'sql' else lang.title()
                break
        
        # Extract course title with better heuristics
        lines = content.split('\n')
        for line in lines[:15]:  # Check first 15 lines
            line = line.strip()
            # Look for title-like lines
            if (len(line) > 10 and len(line) < 150 and 
                not line.startswith('#') and not line.startswith('*') and
                not line.startswith('-') and line[0].isupper()):
                # Additional checks for title-like content
                if not any(char in line for char in ['()', '[]', '{}']):
                    metadata["title"] = line[:100]
                    break
        
        # If no title found, use filename
        if "title" not in metadata:
            clean_filename = file_path.stem.replace('_', ' ').replace('-', ' ').title()
            metadata["title"] = clean_filename[:100]
        
        return metadata

    def create_optimized_chunks(self, documents):
        """Create optimized chunks for educational content with format-aware splitting"""
        # Use different chunk sizes based on document type
        def get_chunk_params(doc):
            file_format = doc.metadata.get('file_format', 'TXT')
            
            if file_format == 'PDF':
                return {'chunk_size': 1000, 'chunk_overlap': 150}  # PDFs might have formatting issues
            elif file_format == 'DOCX':
                return {'chunk_size': 1200, 'chunk_overlap': 150}  # Word docs are usually well-structured
            elif file_format == 'MD':
                return {'chunk_size': 1100, 'chunk_overlap': 100}  # Markdown has natural break points
            else:  # TXT
                return {'chunk_size': 1200, 'chunk_overlap': 150}  # Default for text files
        
        logger.info("Creating optimized document chunks...")
        all_chunks = []
        
        for doc in documents:
            params = get_chunk_params(doc)
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=params['chunk_size'],
                chunk_overlap=params['chunk_overlap'],
                length_function=len,
                separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""],
                is_separator_regex=False
            )
            
            chunks = text_splitter.split_documents([doc])
            
            # Filter and enhance chunks
            valid_chunks = []
            for chunk in chunks:
                content = chunk.page_content.strip()
                
                # More lenient minimum length for different formats
                min_length = 100 if doc.metadata.get('file_format') == 'PDF' else 150
                
                if len(content) < min_length:
                    continue
                
                # Enhanced chunk metadata
                chunk.metadata.update({
                    "chunk_id": f"{doc.metadata['source']}_{len(valid_chunks)+1}",
                    "chunk_length": len(content),
                    "parent_course": doc.metadata.get("course_title", "Unknown"),
                    "source_format": doc.metadata.get('file_format', 'UNKNOWN'),
                    "chunk_params": f"size_{params['chunk_size']}_overlap_{params['chunk_overlap']}"
                })
                
                valid_chunks.append(chunk)
            
            all_chunks.extend(valid_chunks)
            logger.info(f"Created {len(valid_chunks)} chunks from {doc.metadata['source']} ({doc.metadata.get('file_format', 'UNKNOWN')})")
        
        # Log chunking statistics by format
        chunk_stats = {}
        for chunk in all_chunks:
            fmt = chunk.metadata.get('source_format', 'UNKNOWN')
            chunk_stats[fmt] = chunk_stats.get(fmt, 0) + 1
        
        logger.info("Chunk distribution by format:")
        for fmt, count in chunk_stats.items():
            logger.info(f"  • {fmt}: {count} chunks")
        
        logger.info(f"Total chunks created: {len(all_chunks)}")
        return all_chunks

    def batch_upload_with_monitoring(self, chunks):
        """Upload chunks with comprehensive monitoring and error recovery"""
        if not chunks:
            logger.error("No chunks to upload")
            return False
        
        self.vector_store = PineconeVectorStore(index=self.index, embedding=self.embeddings)
        
        batch_size = 40  # Optimal batch size for stability
        total_batches = (len(chunks) - 1) // batch_size + 1
        successful_uploads = 0
        
        logger.info(f"Uploading {len(chunks)} chunks in {total_batches} batches...")
        
        for batch_num in range(total_batches):
            start_idx = batch_num * batch_size
            end_idx = min(start_idx + batch_size, len(chunks))
            batch_chunks = chunks[start_idx:end_idx]
            
            # Generate unique IDs for this batch
            batch_ids = [f"chunk_{start_idx + i + 1:06d}" for i in range(len(batch_chunks))]
            
            retry_count = 0
            max_retries = 3
            
            while retry_count < max_retries:
                try:
                    self.vector_store.add_documents(
                        documents=batch_chunks,
                        ids=batch_ids
                    )
                    
                    successful_uploads += len(batch_chunks)
                    logger.info(f"✅ Batch {batch_num + 1}/{total_batches} uploaded successfully ({len(batch_chunks)} chunks)")
                    break
                    
                except Exception as e:
                    retry_count += 1
                    wait_time = min(2 ** retry_count, 10)
                    
                    if retry_count < max_retries:
                        logger.warning(f"Batch {batch_num + 1} failed, retrying in {wait_time}s: {e}")
                        time.sleep(wait_time)
                    else:
                        logger.error(f"❌ Batch {batch_num + 1} failed permanently: {e}")
                        break
            
            # Rate limiting between batches
            if batch_num < total_batches - 1:
                time.sleep(1)
        
        logger.info(f"Upload completed: {successful_uploads}/{len(chunks)} chunks uploaded")
        return successful_uploads == len(chunks)

    def verify_indexing(self, expected_count, max_wait_time=120):
        """Verify indexing completion with timeout"""
        logger.info("Verifying index status...")
        start_time = time.time()
        
        while time.time() - start_time < max_wait_time:
            try:
                stats = self.index.describe_index_stats()
                current_count = stats.total_vector_count
                
                logger.info(f"Index contains {current_count}/{expected_count} vectors")
                
                if current_count >= expected_count * 0.95:  # Allow 5% margin
                    logger.info("✅ Indexing verification successful")
                    return True
                
                time.sleep(10)
                
            except Exception as e:
                logger.warning(f"Index verification error: {e}")
                time.sleep(5)
        
        logger.warning("⚠️ Index verification timed out")
        return False

    def run_ingestion_pipeline(self):
        """Execute complete ingestion pipeline with comprehensive monitoring"""
        logger.info("🚀 Starting multi-format content ingestion pipeline...")
        logger.info(f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}")
        
        # Initialize components
        if not self.initialize_pinecone():
            logger.error("❌ Pipeline failed: Pinecone initialization")
            return False
            
        if not self.initialize_embeddings():
            logger.error("❌ Pipeline failed: Embeddings initialization")
            return False
        
        # Load and process documents
        documents = self.load_course_documents()
        if not documents:
            logger.error("❌ Pipeline failed: No documents loaded")
            return False
        
        # Create optimized chunks
        chunks = self.create_optimized_chunks(documents)
        if not chunks:
            logger.error("❌ Pipeline failed: No valid chunks created")
            return False
        
        # Upload to vector store
        upload_success = self.batch_upload_with_monitoring(chunks)
        if not upload_success:
            logger.error("❌ Pipeline failed: Upload incomplete")
            return False
        
        # Verify indexing
        indexing_verified = self.verify_indexing(len(chunks))
        
        # Final report
        try:
            final_stats = self.index.describe_index_stats()
            logger.info("🎉 Ingestion pipeline completed successfully!")
            logger.info("📊 Final Statistics:")
            logger.info(f"  • Documents processed: {len(documents)}")
            logger.info(f"  • Chunks created: {len(chunks)}")
            logger.info(f"  • Vectors indexed: {final_stats.total_vector_count}")
            logger.info(f"  • Index verification: {'✅ Passed' if indexing_verified else '⚠️ Partial'}")
            
            return True
            
        except Exception as e:
            logger.error(f"Final statistics error: {e}")
            return upload_success

def main():
    """Main execution function"""
    try:
        ingestion = MultiFormatContentIngestion()
        success = ingestion.run_ingestion_pipeline()
        
        if success:
            logger.info("🎯 Ingestion completed successfully - Ready for production!")
        else:
            logger.error("💥 Ingestion failed - Check logs for details")
            exit(1)
            
    except KeyboardInterrupt:
        logger.info("⏸️ Ingestion interrupted by user")
        exit(0)
    except Exception as e:
        logger.error(f"💥 Unexpected error: {e}")
        exit(1)

if __name__ == "__main__":
    main()